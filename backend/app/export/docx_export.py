"""Export a full analysis report as a .docx file using python-docx."""
import json
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def _add_list(doc, items):
    for item in items:
        doc.add_paragraph(str(item), style="List Bullet")


def export_paper_to_docx(paper, output_path: Path) -> Path:
    doc = Document()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(paper.title or paper.filename)
    run.bold = True
    run.font.size = Pt(20)

    authors = json.loads(paper.authors) if paper.authors else []
    if authors:
        author_p = doc.add_paragraph(", ".join(authors))
        author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Difficulty: {paper.difficulty_label or 'N/A'} "
                       f"({paper.difficulty_score or 0}/100)  |  "
                       f"Est. reading time: {paper.estimated_reading_minutes or 'N/A'} min")

    _add_heading(doc, "Abstract", level=2)
    doc.add_paragraph(paper.abstract or "Not available")

    _add_heading(doc, "Executive Summary", level=2)
    doc.add_paragraph(paper.executive_summary or "")

    _add_heading(doc, "Standard Summary", level=2)
    doc.add_paragraph(paper.standard_summary or "")

    _add_heading(doc, "Detailed Summary", level=2)
    doc.add_paragraph(paper.detailed_summary or "")

    _add_heading(doc, "Simplified Technical Explanation", level=2)
    doc.add_paragraph(paper.simplified_explanation or "")

    _add_heading(doc, "Key Contributions", level=2)
    _add_list(doc, json.loads(paper.key_contributions or "[]"))

    _add_heading(doc, "Research Gaps", level=2)
    _add_list(doc, json.loads(paper.research_gaps or "[]"))

    _add_heading(doc, "Future Scope", level=2)
    _add_list(doc, json.loads(paper.future_scope or "[]"))

    _add_heading(doc, "Methodology Breakdown", level=2)
    methodology = json.loads(paper.methodology_json or "{}")
    doc.add_paragraph("Algorithms: " + ", ".join(methodology.get("algorithms", [])) or "N/A")
    doc.add_paragraph("Datasets: " + ", ".join(methodology.get("datasets", [])) or "N/A")
    doc.add_paragraph("Evaluation Metrics: " + ", ".join(methodology.get("evaluation_metrics", [])) or "N/A")
    if methodology.get("pipeline_steps"):
        doc.add_paragraph("Pipeline Steps:")
        _add_list(doc, methodology["pipeline_steps"])

    _add_heading(doc, "Keywords", level=2)
    keywords = json.loads(paper.keywords_json or "[]")
    doc.add_paragraph(", ".join(k["keyword"] for k in keywords))

    _add_heading(doc, "Viva Questions", level=2)
    viva = json.loads(paper.viva_questions_json or "{}")
    for level_name, questions in viva.items():
        doc.add_heading(level_name.replace("_", " ").title(), level=3)
        _add_list(doc, questions)

    _add_heading(doc, "Citations", level=2)
    citations = json.loads(paper.citations_json or "{}")
    for style, citation in citations.items():
        doc.add_paragraph(f"{style.upper()}: {citation}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
